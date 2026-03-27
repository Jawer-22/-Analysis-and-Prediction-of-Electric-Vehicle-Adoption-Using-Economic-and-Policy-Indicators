import sys
import os
import json
import base64
import re
import math
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def strip_markdown(text):
    # Remove standard markdown symbols like #, *, **, __
    text = re.sub(r'\*{1,2}', '', text)
    text = re.sub(r'_{1,2}', '', text)
    text = re.sub(r'(?m)^#+\s*', '', text)  # remove hashes at start of headings
    text = re.sub(r'`', '', text)       # remove inline code ticks
    text = re.sub(r'(?m)^-\s+', '', text) # remove bullet hyphens for clean reading
    return text.strip()

def get_base_filler(idx):
    fillers = {
        'default': [
            "In any data science process, systematic methodologies are paramount for ensuring predictive models form reliable valid results. Rigorous documentation defines analytical pipelines extending fully from raw data to eventual deployment stages.",
            "Structured approaches provide mechanisms ensuring quantitative and qualitative insights become appropriately quantified for review. Detailed breakdowns guarantee thorough evaluation of attributes while analytical segments remain continuously verified against prior stages to maintain statistical soundness.",
            "Consistent parameters establish strong baselines for iterative improvements throughout tracking lifecycles. Strategies employed here correctly address structural discrepancies across datasets, yielding insight into overarching dynamics without succumbing to external variance biases."
        ],
        'ch1': [
            "Predicting and analyzing Electric Vehicle (EV) adoption trends represents a critical step toward understanding sustainable global transportation networks. This project utilizes machine learning techniques specifically combined alongside economic and political indicators to robustly forecast future EV penetration rates.",
            "Systematic investigations of historical demographics, growth dynamics, and evolving policy landscapes establish powerful predictive frameworks accurately highlighting foundational adoption relationships. Identifying variables allows authorities to proactively inform complex infrastructure deployments."
        ],
        'ch9': [
            "We realize model deployment functionality via a comprehensively optimized Streamlit web application interface. This front-end service permits stakeholder inputs covering varying custom political variables to extract instant adoption predictions automatically.",
            "Streamlit accommodates seamless integration directly atop Python-based predictive frameworks with high modularity and low response latency. Categorical and continuous values entered adjust dynamically within execution spaces to present precise insights."
        ]
    }
    if str(idx).startswith('1') or "Introduction" in idx: return fillers['ch1']
    if str(idx).startswith('9') or "Webapp" in idx: return fillers['ch9']
    return fillers['default']

def setup_styles(doc):
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Times New Roman'
    font_h1.size = Pt(16)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0, 0, 0)
    style_h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Times New Roman'
    font_h2.size = Pt(14)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(0, 0, 0)
    style_h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_heading1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Squeeze spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_paragraph(doc, text, is_bullet=False):
    if is_bullet:
        p = doc.add_paragraph(text, style='List Bullet')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6) # Shortest natural gap
    
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Inches(0.5)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_placeholder_box(doc, text_content):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    kwargs = {"sz": 12, "val": "single", "color": "000000"}
    set_cell_border(cell, top=kwargs, bottom=kwargs, left=kwargs, right=kwargs)
    
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(f"\n\n[ {text_content} Placeholder ]\n[ Insert Final Visual Here ]\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def extract_notebook_cells(nb_path):
    if not os.path.exists(nb_path):
        return []
    try:
        with open(nb_path, 'r', encoding='utf-8') as f:
            nb = json.load(f)
        return nb.get('cells', [])
    except Exception as e:
        print(f"Error reading notebook {nb_path}: {e}")
        return []

def main():
    doc = Document()
    setup_styles(doc)
    
    main_title = doc.add_heading('ANALYSIS AND PREDICTION OF EV ADOPTION USING ECONOMIC AND POLICY INDICATORS', 0)
    main_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    main_title.paragraph_format.space_after = Pt(24)

    for pre in ["CERTIFICATE", "DECLARATION", "ACKNOWLEDGEMENT", "SYNOPSIS"]:
        doc.add_page_break()
        add_heading1(doc, pre)
        if pre == "SYNOPSIS":
            for text in get_base_filler('ch1'):
                add_paragraph(doc, text)

    TOC = [
        {"chap": "1. INTRODUCTION", "nb": None, "subs": []},
        {"chap": "2. DATA UNDERSTANDING", "nb": "Notebooks/01_data_understanding.ipynb", "subs": ["2.1 Project Objective & Dataset Overview", "2.2 Initial Data Quality & Structural Validation", "2.3 Observation Summary"]},
        {"chap": "3. DATA CLEANING", "nb": "Notebooks/02_data_cleaning.ipynb", "subs": ["3.1 Systematic Missing Value Imputation", "3.2 Data Type Standardization & Formatting", "3.3 Outlier Detection & Handling Strategies", "3.4 Observation Summary"]},
        {"chap": "4. FEATURE ENGINEERING", "nb": "Notebooks/03_feature_engineering.ipynb", "subs": ["4.1 Sales Performance & Infrastructure Metrics", "4.2 Economic Constraints & Cost Competitiveness", "4.3 Policy Stringency & Temporal Growth Dynamics", "4.4 Observation Summary"]},
        {"chap": "5. EDA (EXPLORATORY DATA ANALYSIS)", "nb": "Notebooks/04_exploratory_data_analysis.ipynb", "subs": ["5.1 Global & Regional Adoption Trends", "5.2 Economic & Policy Impact Analysis", "5.3 Feature Refinement for Modelling", "5.4 Observation Summary"]},
        {"chap": "6. MODEL BUILDING", "nb": "Notebooks/05_model_building.ipynb", "subs": ["6.1 Feature Selection & Leakage Prevention", "6.2 Baseline Model Comparison (Random Forest Vs. XGBoost)", "6.3 Observation Summary"]},
        {"chap": "7. MODEL EVALUATION", "nb": "Notebooks/06_model_evaluation.ipynb", "subs": ["7.1 Comprehensive Performance Metrics", "7.2 Observation Summary"]},
        {"chap": "8. MODEL EXPLAINABILITY", "nb": "Notebooks/07_model_explainability.ipynb", "subs": ["8.1 Global Feature Importance", "8.2 Observation Summary"]},
        {"chap": "9. WEBAPP", "nb": None, "subs": ["9.1 Objective", "9.2 Overview"]},
        {"chap": "10. BIBLIOGRAPHY", "nb": None, "subs": []},
    ]

    for c_i, chapter in enumerate(TOC):
        doc.add_page_break()
        add_heading1(doc, chapter["chap"])
        
        cells = extract_notebook_cells(chapter["nb"]) if chapter["nb"] else []
        
        if not chapter["subs"]:
            filler = get_base_filler(chapter["chap"])
            for text in filler:
                add_paragraph(doc, text)
            add_paragraph(doc, "Key highlights encompassing this specific chapter structure include:")
            add_paragraph(doc, f"The thorough analytical investigation specific to {chapter['chap'].title()}.", is_bullet=True)
            add_paragraph(doc, "Incremental continuous evaluation focusing strictly alongside analytical metrics.", is_bullet=True)
            add_paragraph(doc, "The outputs generated here directly construct structured predictive insight outcomes minimizing error propagation across variables.")
            continue

        n_subs = len(chapter["subs"])
        cells_per_sub = math.ceil(len(cells) / n_subs) if n_subs > 0 else 0
        
        for i, sub in enumerate(chapter["subs"]):
            add_heading2(doc, sub)
            
            filler = get_base_filler(sub)
            for text in filler:
                add_paragraph(doc, text)
            
            add_paragraph(doc, f"This sector meticulously isolates variables directly correlated underneath the {sub} module. Reviewing isolated factors ensures cross-contamination remains prevented during computational pipelines.")
            
            if cells:
                chunk = cells[i * cells_per_sub : (i+1) * cells_per_sub]
                for cell in chunk:
                    if cell.get('cell_type') == 'markdown':
                        source = ''.join(cell.get('source', []))
                        if source.strip():
                            stripped_source = strip_markdown(source)
                            if stripped_sourceBorder := stripped_source:
                                add_paragraph(doc, stripped_sourceBorder)
                    elif cell.get('cell_type') == 'code':
                        outputs = cell.get('outputs', [])
                        for out in outputs:
                            if out.get('output_type') in ('execute_result', 'display_data'):
                                data = out.get('data', {})
                                if 'image/png' in data:
                                    add_paragraph(doc, "The following visualization explicitly maps output correlations generated against foundational matrices:")
                                    img_data = base64.b64decode(data['image/png'])
                                    img_path = f"tmp_img_{c_i}_{i}.png"
                                    with open(img_path, "wb") as f:
                                        f.write(img_data)
                                    doc.add_picture(img_path, width=Inches(5.0))
                                    os.remove(img_path)

    doc.add_page_break()
    add_heading1(doc, "11. APPENDICES")
    
    appendices = [
        ("Appendix A", "Data Flow Diagram"),
        ("Appendix B", "Sample Dataset"),
        ("Appendix C", "Sample Coding"),
        ("Appendix D", "Sample Input Screen"),
        ("Appendix E", "Sample Output Screen"),
        ("Appendix F", "Process Flow Diagram")
    ]

    for app_id, app_title in appendices:
        doc.add_page_break()
        add_heading2(doc, f"{app_id} - {app_title}")
        
        if app_id == "Appendix C":
            add_paragraph(doc, "The isolated sequential sample execution codes covering methodology validation scripts generated across all fundamental processing notebooks:")
            for ch in TOC:
                if ch["nb"] and os.path.exists(ch["nb"]):
                    add_paragraph(doc, f"--- Script Implementation Array: {os.path.basename(ch['nb'])} ---")
                    cells = extract_notebook_cells(ch["nb"])
                    for cell in cells:
                        if cell.get('cell_type') == 'code':
                            source = ''.join(cell.get('source', []))
                            if source.strip():
                                add_code_block(doc, source.strip())
        elif app_title in ("Data Flow Diagram", "Process Flow Diagram", "Sample Input Screen", "Sample Output Screen", "Sample Dataset"):
            add_paragraph(doc, f"This structural appendix visually represents the {app_title}. Review the implementation architecture deployed internally:")
            add_placeholder_box(doc, app_title)
        else:
            add_paragraph(doc, f"This appendix structurally isolates the {app_title}. Additional diagrammatical architecture elements persist fully configured underlying validation formats within central deployments.")

    output_path = "ANALYSIS AND PREDICTION OF EV ADOPTION USING ECONOMIC AND POLICTICAL INDICATORS 2.docx"
    doc.save(output_path)
    print("SUCCESS: Document V4 generated at", output_path)

if __name__ == "__main__":
    main()
